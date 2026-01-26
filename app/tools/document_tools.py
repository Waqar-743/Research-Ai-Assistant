"""
Document Processing Tools
PDF text extraction, metadata parsing, citation detection, and document utilities.
"""

import io
import re
from typing import Dict, Any, List, Optional, Tuple
from datetime import datetime

from app.utils.logging import logger

# Lazy imports for optional dependencies
pdfplumber = None
docx = None


def _get_pdfplumber():
    """Lazy load pdfplumber."""
    global pdfplumber
    if pdfplumber is None:
        try:
            import pdfplumber as _pdfplumber
            pdfplumber = _pdfplumber
        except ImportError:
            raise ImportError("pdfplumber is required for PDF processing. Install with: pip install pdfplumber")
    return pdfplumber


def _get_docx():
    """Lazy load python-docx."""
    global docx
    if docx is None:
        try:
            import docx as _docx
            docx = _docx
        except ImportError:
            raise ImportError("python-docx is required for DOCX processing. Install with: pip install python-docx")
    return docx


class DocumentTools:
    """Tools for document processing and text extraction."""
    
    # Supported MIME types
    SUPPORTED_TYPES = {
        "application/pdf": "pdf",
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document": "docx",
        "text/plain": "txt",
        "text/markdown": "md"
    }
    
    # Max file size (5MB)
    MAX_FILE_SIZE = 5 * 1024 * 1024
    
    # Max files per batch
    MAX_BATCH_SIZE = 10
    
    def __init__(self):
        """Initialize document tools."""
        pass
    
    @staticmethod
    def validate_file(
        filename: str,
        content_type: str,
        file_size: int
    ) -> Tuple[bool, Optional[str], Optional[str]]:
        """
        Validate an uploaded file.
        
        Returns:
            Tuple of (is_valid, error_message, document_type)
        """
        # Check file size
        if file_size > DocumentTools.MAX_FILE_SIZE:
            return False, f"File size exceeds maximum allowed ({DocumentTools.MAX_FILE_SIZE / 1024 / 1024:.1f}MB)", None
        
        if file_size == 0:
            return False, "File is empty", None
        
        # Check content type
        if content_type not in DocumentTools.SUPPORTED_TYPES:
            # Try to infer from filename
            ext = filename.lower().split('.')[-1] if '.' in filename else ''
            type_map = {"pdf": "pdf", "docx": "docx", "txt": "txt", "md": "md"}
            if ext in type_map:
                return True, None, type_map[ext]
            return False, f"Unsupported file type: {content_type}. Supported: PDF, DOCX, TXT, MD", None
        
        return True, None, DocumentTools.SUPPORTED_TYPES[content_type]
    
    async def extract_text(
        self,
        file_content: bytes,
        document_type: str,
        filename: str
    ) -> Dict[str, Any]:
        """
        Extract text and metadata from a document.
        
        Returns:
            Dict with extracted_text, page_count, word_count, metadata
        """
        try:
            if document_type == "pdf":
                return await self._extract_pdf(file_content)
            elif document_type == "docx":
                return await self._extract_docx(file_content)
            elif document_type in ("txt", "md"):
                return await self._extract_text_file(file_content, filename)
            else:
                raise ValueError(f"Unsupported document type: {document_type}")
        except Exception as e:
            logger.error(f"Text extraction failed for {filename}: {e}")
            return {
                "extracted_text": None,
                "page_count": None,
                "word_count": None,
                "metadata": {},
                "error": str(e)
            }
    
    async def _extract_pdf(self, file_content: bytes) -> Dict[str, Any]:
        """Extract text from PDF using pdfplumber."""
        pdf_lib = _get_pdfplumber()
        
        text_parts = []
        metadata = {}
        page_count = 0
        
        with pdf_lib.open(io.BytesIO(file_content)) as pdf:
            page_count = len(pdf.pages)
            metadata = pdf.metadata or {}
            
            for page in pdf.pages:
                page_text = page.extract_text()
                if page_text:
                    text_parts.append(page_text)
        
        full_text = "\n\n".join(text_parts)
        word_count = len(full_text.split()) if full_text else 0
        
        return {
            "extracted_text": full_text,
            "page_count": page_count,
            "word_count": word_count,
            "metadata": {
                "title": metadata.get("Title", ""),
                "author": metadata.get("Author", ""),
                "subject": metadata.get("Subject", ""),
                "creator": metadata.get("Creator", ""),
                "creation_date": metadata.get("CreationDate", ""),
                "modification_date": metadata.get("ModDate", "")
            }
        }
    
    async def _extract_docx(self, file_content: bytes) -> Dict[str, Any]:
        """Extract text from DOCX using python-docx."""
        docx_lib = _get_docx()
        
        doc = docx_lib.Document(io.BytesIO(file_content))
        
        text_parts = []
        for para in doc.paragraphs:
            if para.text.strip():
                text_parts.append(para.text)
        
        # Also extract from tables
        for table in doc.tables:
            for row in table.rows:
                row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
                if row_text:
                    text_parts.append(row_text)
        
        full_text = "\n\n".join(text_parts)
        word_count = len(full_text.split()) if full_text else 0
        
        # Extract core properties
        core_props = doc.core_properties
        metadata = {
            "title": core_props.title or "",
            "author": core_props.author or "",
            "subject": core_props.subject or "",
            "created": str(core_props.created) if core_props.created else "",
            "modified": str(core_props.modified) if core_props.modified else ""
        }
        
        return {
            "extracted_text": full_text,
            "page_count": None,  # DOCX doesn't have fixed pages
            "word_count": word_count,
            "metadata": metadata
        }
    
    async def _extract_text_file(self, file_content: bytes, filename: str) -> Dict[str, Any]:
        """Extract text from TXT or MD file."""
        # Try different encodings
        encodings = ['utf-8', 'utf-16', 'latin-1', 'cp1252']
        text = None
        
        for encoding in encodings:
            try:
                text = file_content.decode(encoding)
                break
            except (UnicodeDecodeError, LookupError):
                continue
        
        if text is None:
            text = file_content.decode('utf-8', errors='replace')
        
        word_count = len(text.split()) if text else 0
        
        return {
            "extracted_text": text,
            "page_count": None,
            "word_count": word_count,
            "metadata": {
                "filename": filename,
                "encoding": "utf-8"
            }
        }
    
    def extract_citations(self, text: str) -> List[Dict[str, Any]]:
        """
        Extract potential citations from text.
        Uses pattern matching to identify common citation formats.
        """
        citations = []
        
        if not text:
            return citations
        
        # Pattern for parenthetical citations: (Author, Year) or (Author et al., Year)
        paren_pattern = r'\(([A-Z][a-zA-Z\-]+(?:\s+(?:et\s+al\.|&\s+[A-Z][a-zA-Z\-]+))?),?\s*(\d{4}[a-z]?)\)'
        for match in re.finditer(paren_pattern, text):
            citations.append({
                "raw_text": match.group(0),
                "authors": [match.group(1)],
                "year": int(match.group(2)[:4]),
                "style": "parenthetical",
                "position": match.start()
            })
        
        # Pattern for reference list items (basic APA-like format)
        ref_pattern = r'([A-Z][a-zA-Z\-]+(?:,\s+[A-Z]\.(?:\s*[A-Z]\.)?)?(?:,?\s+(?:&|and)\s+[A-Z][a-zA-Z\-]+(?:,\s+[A-Z]\.(?:\s*[A-Z]\.)?)?)*)\s*\((\d{4}[a-z]?)\)\.\s*([^.]+)\.'
        for match in re.finditer(ref_pattern, text):
            citations.append({
                "raw_text": match.group(0),
                "authors": [a.strip() for a in re.split(r',\s*(?:&|and)\s*', match.group(1))],
                "year": int(match.group(2)[:4]),
                "title": match.group(3).strip(),
                "style": "reference",
                "position": match.start()
            })
        
        # Pattern for DOIs
        doi_pattern = r'(?:doi:|https?://doi\.org/)?(10\.\d{4,}/[^\s]+)'
        for match in re.finditer(doi_pattern, text, re.IGNORECASE):
            # Check if this DOI is already part of an existing citation
            doi = match.group(1).rstrip('.')
            existing = next((c for c in citations if c.get("position", 0) <= match.start() <= c.get("position", 0) + 500), None)
            if existing:
                existing["doi"] = doi
            else:
                citations.append({
                    "raw_text": match.group(0),
                    "doi": doi,
                    "style": "doi",
                    "position": match.start()
                })
        
        # Remove duplicates based on position proximity
        unique_citations = []
        for cit in sorted(citations, key=lambda x: x.get("position", 0)):
            if not unique_citations or abs(cit.get("position", 0) - unique_citations[-1].get("position", 0)) > 10:
                unique_citations.append(cit)
        
        return unique_citations
    
    def format_citation(
        self,
        citation: Dict[str, Any],
        style: str = "APA"
    ) -> str:
        """
        Format a citation in the specified style.
        
        Args:
            citation: Citation data with authors, year, title, etc.
            style: Citation style (APA, MLA, Chicago, Harvard)
        """
        authors = citation.get("authors", [])
        year = citation.get("year", "n.d.")
        title = citation.get("title", "")
        publication = citation.get("publication", "")
        doi = citation.get("doi", "")
        url = citation.get("url", "")
        
        if not authors:
            authors = ["Unknown"]
        
        if style.upper() == "APA":
            # APA 7th edition format
            author_str = self._format_authors_apa(authors)
            citation_str = f"{author_str} ({year}). {title}."
            if publication:
                citation_str += f" {publication}."
            if doi:
                citation_str += f" https://doi.org/{doi}"
            elif url:
                citation_str += f" Retrieved from {url}"
            return citation_str
        
        elif style.upper() == "MLA":
            # MLA 9th edition format
            author_str = self._format_authors_mla(authors)
            citation_str = f'{author_str} "{title}."'
            if publication:
                citation_str += f" {publication},"
            citation_str += f" {year}."
            if url:
                citation_str += f" {url}."
            return citation_str
        
        elif style.upper() == "CHICAGO":
            # Chicago 17th edition format
            author_str = self._format_authors_chicago(authors)
            citation_str = f'{author_str} "{title}."'
            if publication:
                citation_str += f" {publication}"
            citation_str += f" ({year})."
            if url:
                citation_str += f" {url}."
            return citation_str
        
        elif style.upper() == "HARVARD":
            # Harvard format
            author_str = self._format_authors_harvard(authors)
            citation_str = f"{author_str} ({year}) '{title}',"
            if publication:
                citation_str += f" {publication}."
            if doi:
                citation_str += f" doi: {doi}"
            elif url:
                citation_str += f" Available at: {url}"
            return citation_str
        
        else:
            # Default to raw format
            return citation.get("raw_text", str(citation))
    
    def _format_authors_apa(self, authors: List[str]) -> str:
        """Format authors in APA style."""
        if len(authors) == 1:
            return authors[0]
        elif len(authors) == 2:
            return f"{authors[0]} & {authors[1]}"
        elif len(authors) <= 20:
            return ", ".join(authors[:-1]) + f", & {authors[-1]}"
        else:
            return ", ".join(authors[:19]) + f", ... {authors[-1]}"
    
    def _format_authors_mla(self, authors: List[str]) -> str:
        """Format authors in MLA style."""
        if len(authors) == 1:
            return authors[0]
        elif len(authors) == 2:
            return f"{authors[0]} and {authors[1]}"
        else:
            return f"{authors[0]}, et al."
    
    def _format_authors_chicago(self, authors: List[str]) -> str:
        """Format authors in Chicago style."""
        if len(authors) == 1:
            return authors[0]
        elif len(authors) == 2:
            return f"{authors[0]} and {authors[1]}"
        elif len(authors) <= 10:
            return ", ".join(authors[:-1]) + f", and {authors[-1]}"
        else:
            return f"{authors[0]} et al."
    
    def _format_authors_harvard(self, authors: List[str]) -> str:
        """Format authors in Harvard style."""
        if len(authors) == 1:
            return authors[0]
        elif len(authors) == 2:
            return f"{authors[0]} and {authors[1]}"
        elif len(authors) == 3:
            return f"{authors[0]}, {authors[1]} and {authors[2]}"
        else:
            return f"{authors[0]} et al."
    
    def get_text_preview(self, text: str, max_length: int = 500) -> str:
        """Get a preview of the text content."""
        if not text:
            return ""
        
        if len(text) <= max_length:
            return text
        
        # Try to break at a sentence boundary
        preview = text[:max_length]
        last_period = preview.rfind('.')
        if last_period > max_length * 0.6:
            return preview[:last_period + 1]
        
        # Break at word boundary
        last_space = preview.rfind(' ')
        if last_space > 0:
            return preview[:last_space] + "..."
        
        return preview + "..."
    
    def calculate_reading_time(self, word_count: int, wpm: int = 200) -> str:
        """Calculate estimated reading time."""
        if not word_count:
            return "< 1 min"
        
        minutes = word_count / wpm
        if minutes < 1:
            return "< 1 min"
        elif minutes < 60:
            return f"{int(minutes)} min"
        else:
            hours = int(minutes / 60)
            remaining_mins = int(minutes % 60)
            return f"{hours}h {remaining_mins}m"

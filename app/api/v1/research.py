"""
Research API Endpoints
Main endpoints for starting and managing research sessions.
Includes hybrid research, document-based chat, and export functionality.
"""

from fastapi import APIRouter, HTTPException, BackgroundTasks, Depends, Query
from fastapi.responses import StreamingResponse, Response
from typing import Optional, List
from datetime import datetime
import uuid
import io
import json

from app.models import (
    ResearchStartRequest,
    ResearchStatusResponse,
    ResearchResultsResponse,
    APIResponse,
    HybridResearchRequest,
    ChatMessageRequest,
    ChatMessageResponse,
    ConversationHistoryResponse,
    ExportRequest,
    ExportResponse
)
from app.database.schemas import ResearchSession, ResearchStatus, ResearchMode
from app.database.repositories import ResearchRepository
from app.database.document_schemas import ConversationMessage, ConversationRole
from app.database.document_repository import (
    DocumentRepository,
    ConversationRepository,
    SettingsRepository
)
from app.services.research_service import ResearchService
from app.agents.document_analyzer import DocumentAnalyzer
from app.utils.logging import logger


router = APIRouter()

# Research service singleton
_research_service: Optional[ResearchService] = None


def get_research_service() -> ResearchService:
    """Get research service instance."""
    global _research_service
    if _research_service is None:
        _research_service = ResearchService()
    return _research_service


@router.post("/start", response_model=APIResponse)
async def start_research(
    request: ResearchStartRequest,
    background_tasks: BackgroundTasks,
    research_service: ResearchService = Depends(get_research_service)
):
    """
    Start a new research session.
    
    This endpoint initiates the multi-agent research workflow.
    Progress updates are delivered via WebSocket connection.
    """
    try:
        # Generate session ID
        session_id = str(uuid.uuid4())
        
        logger.info(f"Starting research session {session_id}: {request.query}")
        
        # Create session in database
        session = ResearchSession(
            research_id=session_id,
            user_id=request.user_id or "anonymous",
            query=request.query,
            status=ResearchStatus.INITIALIZED,
            research_mode=ResearchMode(request.research_mode or "auto"),
            focus_areas=request.focus_areas or [],
            source_preferences=request.source_preferences or [],
            max_sources=request.max_sources or 300,
            report_format=request.report_format or "markdown",
            citation_style=request.citation_style or "APA",
            created_at=datetime.utcnow()
        )
        await session.insert()
        
        # Start research in background
        background_tasks.add_task(
            research_service.execute_research,
            session_id=session_id,
            query=request.query,
            focus_areas=request.focus_areas,
            source_preferences=request.source_preferences,
            max_sources=request.max_sources or 300,
            research_mode=request.research_mode or "auto",
            report_format=request.report_format or "markdown",
            citation_style=request.citation_style or "APA"
        )
        
        return APIResponse(
            status=200,
            message="Research session started successfully",
            data={
                "session_id": session_id,
                "status": "queued",
                "query": request.query,
                "websocket_url": f"/ws/{session_id}"
            }
        )
        
    except Exception as e:
        logger.error(f"Failed to start research: {e}")
        raise HTTPException(
            status_code=500,
            detail=f"Failed to start research: {str(e)}"
        )


@router.get("/{session_id}", response_model=APIResponse)
async def get_research_status(session_id: str):
    """
    Get the current status of a research session.
    """
    try:
        session = await ResearchRepository.get_by_session_id(session_id)
        
        if not session:
            raise HTTPException(
                status_code=404,
                detail=f"Research session {session_id} not found"
            )
        
        response_data = ResearchStatusResponse(
            session_id=session.session_id,
            status=session.status.value,
            progress=session.progress or 0,
            current_phase=session.current_phase,
            query=session.query,
            created_at=session.created_at,
            agent_statuses=session.agent_statuses or {}
        )
        
        return APIResponse(
            status=200,
            message="Status retrieved successfully",
            data=response_data.model_dump()
        )
        
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Failed to get research status: {e}")
        raise HTTPException(
            status_code=500,
            detail=f"Failed to get research status: {str(e)}"
        )


@router.get("/{session_id}/results", response_model=APIResponse)
async def get_research_results(session_id: str):
    """
    Get the complete results of a research session.
    Only available after research is completed.
    """
    try:
        session = await ResearchRepository.get_by_session_id(session_id)
        
        if not session:
            raise HTTPException(
                status_code=404,
                detail=f"Research session {session_id} not found"
            )
        
        if session.status == ResearchStatus.IN_PROGRESS:
            return APIResponse(
                success=False,
                message="Research is still in progress",
                data={
                    "session_id": session_id,
                    "status": session.status.value,
                    "progress": session.progress
                }
            )
        
        if session.status == ResearchStatus.FAILED:
            return APIResponse(
                success=False,
                message="Research failed",
                data={
                    "session_id": session_id,
                    "status": session.status.value,
                    "error": session.error_message
                }
            )
        
        # Build results response
        results = ResearchResultsResponse(
            session_id=session.session_id,
            status=session.status.value,
            query=session.query,
            report=session.final_report or {},
            sources_count=session.sources_count or {},
            findings_count=session.findings_count or 0,
            confidence_summary=session.confidence_summary or {},
            created_at=session.created_at,
            completed_at=session.completed_at
        )
        
        return APIResponse(
            status=200,
            message="Results retrieved successfully",
            data=results.model_dump()
        )
        
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Failed to get research results: {e}")
        raise HTTPException(
            status_code=500,
            detail=f"Failed to get research results: {str(e)}"
        )


@router.post("/{session_id}/cancel", response_model=APIResponse)
async def cancel_research(
    session_id: str,
    research_service: ResearchService = Depends(get_research_service)
):
    """
    Cancel an in-progress research session.
    """
    try:
        session = await ResearchRepository.get_by_session_id(session_id)
        
        if not session:
            raise HTTPException(
                status_code=404,
                detail=f"Research session {session_id} not found"
            )
        
        if session.status != ResearchStatus.IN_PROGRESS:
            return APIResponse(
                success=False,
                message=f"Cannot cancel session with status: {session.status.value}",
                data={"session_id": session_id, "status": session.status.value}
            )
        
        # Cancel the research
        await research_service.cancel_research(session_id)
        
        # Update session status
        session.status = ResearchStatus.CANCELLED
        session.updated_at = datetime.utcnow()
        await session.save()
        
        return APIResponse(
            status=200,
            message="Research cancelled successfully",
            data={"session_id": session_id, "status": "cancelled"}
        )
        
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Failed to cancel research: {e}")
        raise HTTPException(
            status_code=500,
            detail=f"Failed to cancel research: {str(e)}"
        )


@router.post("/{session_id}/feedback", response_model=APIResponse)
async def submit_feedback(
    session_id: str,
    feedback: dict,
    research_service: ResearchService = Depends(get_research_service)
):
    """
    Submit user feedback for supervised mode checkpoints.
    """
    try:
        session = await ResearchRepository.get_by_session_id(session_id)
        
        if not session:
            raise HTTPException(
                status_code=404,
                detail=f"Research session {session_id} not found"
            )
        
        # Process feedback
        approved = feedback.get("approved", True)
        user_feedback = feedback.get("feedback", "")
        modifications = feedback.get("modifications", {})
        
        await research_service.process_feedback(
            session_id=session_id,
            approved=approved,
            feedback=user_feedback,
            modifications=modifications
        )
        
        return APIResponse(
            status=200,
            message="Feedback submitted successfully",
            data={"session_id": session_id, "approved": approved}
        )
        
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Failed to submit feedback: {e}")
        raise HTTPException(
            status_code=500,
            detail=f"Failed to submit feedback: {str(e)}"
        )


# ===========================================
# Hybrid Research (Web + Documents)
# ===========================================

@router.post("/start-hybrid", response_model=APIResponse)
async def start_hybrid_research(
    request: HybridResearchRequest,
    background_tasks: BackgroundTasks,
    research_service: ResearchService = Depends(get_research_service)
):
    """
    Start a hybrid research session combining web search with document analysis.
    
    This endpoint:
    1. Analyzes provided documents for context
    2. Conducts web research based on query + document context
    3. Synthesizes findings from both sources
    4. Produces a comprehensive report
    """
    try:
        session_id = str(uuid.uuid4())
        
        logger.info(f"Starting hybrid research {session_id}: {request.query}")
        
        # Validate documents exist and are processed
        documents_context = []
        if request.document_ids:
            for doc_id in request.document_ids:
                doc = await DocumentRepository.get_by_id(doc_id)
                if not doc:
                    raise HTTPException(
                        status_code=404,
                        detail=f"Document {doc_id} not found"
                    )
                documents_context.append({
                    "document_id": doc.document_id,
                    "filename": doc.filename,
                    "summary": doc.summary,
                    "topics": doc.topics,
                    "key_findings": doc.key_findings
                })
        
        # Create session with hybrid mode
        session = ResearchSession(
            session_id=session_id,
            query=request.query,
            status=ResearchStatus.QUEUED,
            research_mode=ResearchMode("auto"),
            focus_areas=request.focus_areas or [],
            source_preferences=request.source_preferences or [],
            max_sources=request.max_sources or 300,
            report_format=request.report_format or "markdown",
            citation_style=request.citation_style or "APA",
            created_at=datetime.utcnow()
        )
        await session.insert()
        
        # Start hybrid research in background
        background_tasks.add_task(
            research_service.execute_research,
            session_id=session_id,
            query=request.query,
            focus_areas=request.focus_areas,
            source_preferences=request.source_preferences,
            max_sources=request.max_sources or 300,
            research_mode="hybrid",
            report_format=request.report_format or "markdown",
            citation_style=request.citation_style or "APA",
            document_context=documents_context,
            include_web_search=request.include_web_search,
            weight_documents=request.weight_documents
        )
        
        return APIResponse(
            status=200,
            message="Hybrid research session started",
            data={
                "session_id": session_id,
                "status": "queued",
                "query": request.query,
                "documents_count": len(documents_context),
                "include_web_search": request.include_web_search,
                "websocket_url": f"/ws/{session_id}"
            }
        )
        
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Failed to start hybrid research: {e}")
        raise HTTPException(
            status_code=500,
            detail=f"Failed to start hybrid research: {str(e)}"
        )


# ===========================================
# Document-Based Chat
# ===========================================

@router.post("/{session_id}/chat", response_model=APIResponse)
async def chat_with_research(
    session_id: str,
    request: ChatMessageRequest,
    user_id: str = Query(default="anonymous")
):
    """
    Chat with a completed research session or documents.
    
    Enables users to ask follow-up questions about:
    - Research findings
    - Specific details from sources
    - Clarifications on the report
    - Related questions
    """
    try:
        # Get or create conversation
        conversation = await ConversationRepository.get_by_session(session_id)
        
        if not conversation:
            # Fetch session context
            session = await ResearchRepository.get_by_session_id(session_id)
            
            if not session:
                raise HTTPException(
                    status_code=404,
                    detail=f"Research session {session_id} not found"
                )
            
            # Build context from session
            context = {
                "query": session.query,
                "report": session.final_report or {},
                "sources_count": session.sources_count or {}
            }
            
            # Add document context if provided
            documents_context = []
            if request.document_ids:
                for doc_id in request.document_ids:
                    doc = await DocumentRepository.get_by_id(doc_id)
                    if doc:
                        documents_context.append({
                            "document_id": doc.document_id,
                            "filename": doc.filename,
                            "summary": doc.summary,
                            "topics": doc.topics,
                            "key_findings": doc.key_findings,
                            "extracted_text": doc.extracted_text[:10000] if doc.extracted_text else None
                        })
            
            # Create conversation
            conversation = await ConversationRepository.create({
                "session_id": session_id,
                "user_id": user_id,
                "document_ids": request.document_ids or [],
                "context": context,
                "messages": []
            })
        
        # Add user message
        user_message = ConversationMessage(
            role=ConversationRole.USER,
            content=request.message,
            timestamp=datetime.utcnow()
        )
        
        # Get AI response
        analyzer = DocumentAnalyzer()
        
        # Build context for the question
        context_text = ""
        if conversation.context:
            if conversation.context.get("report"):
                report = conversation.context["report"]
                context_text += f"Research Query: {conversation.context.get('query', '')}\n"
                context_text += f"Executive Summary: {report.get('executive_summary', '')}\n"
                context_text += f"Key Findings: {report.get('key_findings', '')}\n"
        
        # Add document context
        if request.document_ids:
            for doc_id in request.document_ids:
                doc = await DocumentRepository.get_by_id(doc_id)
                if doc:
                    context_text += f"\n--- Document: {doc.filename} ---\n"
                    context_text += f"Summary: {doc.summary or 'No summary'}\n"
                    if doc.extracted_text:
                        context_text += f"Content: {doc.extracted_text[:5000]}...\n"
        
        # Get response from analyzer
        response = await analyzer.answer_question(
            question=request.message,
            context=context_text,
            conversation_history=[
                {"role": m.role.value, "content": m.content}
                for m in (conversation.messages or [])[-10:]  # Last 10 messages for context
            ]
        )
        
        assistant_message = ConversationMessage(
            role=ConversationRole.ASSISTANT,
            content=response.get("answer", "I couldn't generate a response."),
            timestamp=datetime.utcnow(),
            metadata={
                "confidence": response.get("confidence"),
                "sources_used": response.get("sources_used", [])
            }
        )
        
        # Update conversation with new messages
        messages = conversation.messages or []
        messages.extend([user_message, assistant_message])
        await ConversationRepository.add_message(conversation.conversation_id, user_message)
        await ConversationRepository.add_message(conversation.conversation_id, assistant_message)
        
        return APIResponse(
            status=200,
            message="Chat response generated",
            data=ChatMessageResponse(
                conversation_id=conversation.conversation_id,
                message_id=assistant_message.message_id,
                role="assistant",
                content=assistant_message.content,
                timestamp=assistant_message.timestamp,
                metadata=assistant_message.metadata
            ).model_dump()
        )
        
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Chat failed: {e}")
        raise HTTPException(
            status_code=500,
            detail=f"Failed to process chat: {str(e)}"
        )


@router.get("/{session_id}/chat/history", response_model=APIResponse)
async def get_chat_history(session_id: str, user_id: str = Query(default="anonymous")):
    """Get chat history for a research session."""
    conversation = await ConversationRepository.get_by_session(session_id)
    
    if not conversation:
        return APIResponse(
            status=200,
            message="No conversation history found",
            data={"session_id": session_id, "messages": []}
        )
    
    messages = [
        {
            "message_id": m.message_id,
            "role": m.role.value,
            "content": m.content,
            "timestamp": m.timestamp.isoformat() if m.timestamp else None,
            "metadata": m.metadata
        }
        for m in (conversation.messages or [])
    ]
    
    return APIResponse(
        status=200,
        message=f"Found {len(messages)} messages",
        data=ConversationHistoryResponse(
            conversation_id=conversation.conversation_id,
            session_id=session_id,
            messages=messages,
            created_at=conversation.created_at,
            updated_at=conversation.updated_at
        ).model_dump()
    )


# ===========================================
# Export
# ===========================================

@router.get("/{session_id}/export", response_model=APIResponse)
async def export_research(
    session_id: str,
    format: str = Query(default="markdown", pattern="^(markdown|pdf|docx|json)$"),
    include_sources: bool = Query(default=True),
    include_metadata: bool = Query(default=True)
):
    """
    Export research results in various formats.
    
    Supported formats:
    - markdown: Plain markdown text
    - pdf: PDF document (requires weasyprint)
    - docx: Word document (requires python-docx)
    - json: Raw JSON data
    """
    try:
        session = await ResearchRepository.get_by_session_id(session_id)
        
        if not session:
            raise HTTPException(
                status_code=404,
                detail=f"Research session {session_id} not found"
            )
        
        if session.status != ResearchStatus.COMPLETED:
            raise HTTPException(
                status_code=400,
                detail="Research must be completed before export"
            )
        
        report = session.final_report or {}
        
        # Build export content based on format
        if format == "json":
            content = {
                "session_id": session_id,
                "query": session.query,
                "report": report,
                "metadata": {
                    "created_at": session.created_at.isoformat() if session.created_at else None,
                    "completed_at": session.completed_at.isoformat() if session.completed_at else None,
                    "sources_count": session.sources_count,
                    "findings_count": session.findings_count
                } if include_metadata else None
            }
            
            return Response(
                content=json.dumps(content, indent=2, default=str),
                media_type="application/json",
                headers={
                    "Content-Disposition": f'attachment; filename="research_{session_id}.json"'
                }
            )
        
        elif format == "markdown":
            md_content = _build_markdown_export(session, report, include_sources, include_metadata)
            
            return Response(
                content=md_content,
                media_type="text/markdown",
                headers={
                    "Content-Disposition": f'attachment; filename="research_{session_id}.md"'
                }
            )
        
        elif format == "docx":
            try:
                from docx import Document
                from docx.shared import Pt, Inches
                
                doc = Document()
                doc.add_heading(f"Research Report: {session.query}", 0)
                
                if include_metadata:
                    doc.add_paragraph(f"Generated: {datetime.utcnow().strftime('%Y-%m-%d %H:%M:%S')}")
                    doc.add_paragraph(f"Session ID: {session_id}")
                
                # Executive Summary
                if report.get("executive_summary"):
                    doc.add_heading("Executive Summary", level=1)
                    doc.add_paragraph(report["executive_summary"])
                
                # Key Findings
                if report.get("key_findings"):
                    doc.add_heading("Key Findings", level=1)
                    doc.add_paragraph(report["key_findings"])
                
                # Main Content
                if report.get("detailed_analysis"):
                    doc.add_heading("Detailed Analysis", level=1)
                    doc.add_paragraph(report["detailed_analysis"])
                
                # Sources
                if include_sources and report.get("sources"):
                    doc.add_heading("Sources", level=1)
                    for source in report["sources"][:50]:
                        if isinstance(source, dict):
                            doc.add_paragraph(f"• {source.get('title', 'Untitled')} - {source.get('url', '')}")
                        else:
                            doc.add_paragraph(f"• {source}")
                
                # Save to buffer
                buffer = io.BytesIO()
                doc.save(buffer)
                buffer.seek(0)
                
                return StreamingResponse(
                    buffer,
                    media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    headers={
                        "Content-Disposition": f'attachment; filename="research_{session_id}.docx"'
                    }
                )
                
            except ImportError:
                raise HTTPException(
                    status_code=501,
                    detail="DOCX export not available. Install python-docx package."
                )
        
        elif format == "pdf":
            # PDF export would require additional libraries (weasyprint, etc.)
            raise HTTPException(
                status_code=501,
                detail="PDF export not yet implemented. Use markdown or docx format."
            )
        
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Export failed: {e}")
        raise HTTPException(
            status_code=500,
            detail=f"Failed to export research: {str(e)}"
        )


def _build_markdown_export(session, report: dict, include_sources: bool, include_metadata: bool) -> str:
    """Build markdown content for export."""
    lines = []
    
    lines.append(f"# Research Report: {session.query}\n")
    
    if include_metadata:
        lines.append("---")
        lines.append(f"**Session ID:** {session.session_id}")
        lines.append(f"**Created:** {session.created_at.strftime('%Y-%m-%d %H:%M:%S') if session.created_at else 'N/A'}")
        lines.append(f"**Completed:** {session.completed_at.strftime('%Y-%m-%d %H:%M:%S') if session.completed_at else 'N/A'}")
        if session.sources_count:
            lines.append(f"**Sources:** {session.sources_count}")
        lines.append("---\n")
    
    # Executive Summary
    if report.get("executive_summary"):
        lines.append("## Executive Summary\n")
        lines.append(report["executive_summary"])
        lines.append("")
    
    # Key Findings
    if report.get("key_findings"):
        lines.append("## Key Findings\n")
        lines.append(report["key_findings"])
        lines.append("")
    
    # Detailed Analysis
    if report.get("detailed_analysis"):
        lines.append("## Detailed Analysis\n")
        lines.append(report["detailed_analysis"])
        lines.append("")
    
    # Methodology
    if report.get("methodology"):
        lines.append("## Methodology\n")
        lines.append(report["methodology"])
        lines.append("")
    
    # Recommendations
    if report.get("recommendations"):
        lines.append("## Recommendations\n")
        lines.append(report["recommendations"])
        lines.append("")
    
    # Sources
    if include_sources and report.get("sources"):
        lines.append("## Sources\n")
        for i, source in enumerate(report["sources"][:50], 1):
            if isinstance(source, dict):
                title = source.get("title", "Untitled")
                url = source.get("url", "")
                lines.append(f"{i}. [{title}]({url})")
            else:
                lines.append(f"{i}. {source}")
        lines.append("")
    
    return "\n".join(lines)

